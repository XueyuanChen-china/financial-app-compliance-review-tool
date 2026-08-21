from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from compliance_review.compilation.models import (
    ComplianceSource,
    ControlDraft,
    ControlDraftSet,
    Obligation,
    ObligationSet,
    SourceRegistry,
    SourceSection,
)
from compliance_review.compilation.service import Phase2CompilationError, Phase2CompilationService
from compliance_review.compilation.source_registry import SourceRegistryBuilder
from compliance_review.compilation.validator import (
    ControlValidator,
    validate_applicability_expression,
)
from compliance_review.domain.models import (
    AcceptanceCriterion,
    EvidenceClaim,
    EvidenceProofRoute,
    EvidenceRequirement,
    SourceRef,
)
from compliance_review.review.models import ModelResponse
from compliance_review.review.provider import StaticModelProvider


def _write_text_sources(root: Path) -> list[Path]:
    markdown = root / "google-play.md"
    markdown.write_text(
        "# Personal Loans\n\nPersonal loan terms must be disclosed before approval.\n",
        encoding="utf-8",
    )
    text = root / "secp.txt"
    text.write_text("The lender must maintain a complaint process.\n", encoding="utf-8")
    docx_path = root / "notice.docx"
    document = Document()
    document.add_heading("Disclosure", level=1)
    document.add_paragraph("The borrower must receive clear loan information.")
    document.save(docx_path)
    pdf_path = root / "pricing.pdf"
    _write_simple_pdf(pdf_path)
    return [markdown, text, docx_path, pdf_path]


def _write_simple_pdf(path: Path) -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = writer._add_object(
        DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
    )
    stream = DecodedStreamObject()
    stream.set_data(b"BT /F1 12 Tf 72 720 Td (The loan term must be disclosed.) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)


def test_source_registry_extracts_md_txt_pdf_docx_with_provenance(tmp_path: Path) -> None:
    _write_text_sources(tmp_path)
    registry = SourceRegistryBuilder().build([tmp_path], {tmp_path.as_posix(): "country_regulator"})

    assert len(registry.sources) == 4
    assert {source.media_type for source in registry.sources} == {"md", "txt", "pdf", "docx"}
    assert all(source.sha256 and len(source.sha256) == 64 for source in registry.sources)
    assert all(source.extraction_status == "ok" for source in registry.sources)
    assert all(source.sections for source in registry.sources)
    assert all(source.source_family == "country_regulator" for source in registry.sources)


def _provider_for_valid_compilation() -> StaticModelProvider:
    def response(request: object) -> ModelResponse:
        kind = request.request_kind  # type: ignore[attr-defined]
        payload = json.loads(request.messages[1]["content"])  # type: ignore[attr-defined]
        if kind == "obligation_extraction":
            source = payload["sources"][0]
            section = source["sections"][0]["section_id"]
            source_id = source["source_id"]
            return ModelResponse(
                content=json.dumps(
                    {
                        "contract": "obligation_extraction_batch.v1",
                        "version": "1.0",
                        "source_id": source_id,
                        "batch_id": payload["batch_id"],
                        "section_decisions": [
                            {
                                "section_id": section,
                                "decision": "obligations_extracted",
                                "obligation_ids": ["obl.loan.disclosure"],
                            }
                        ],
                        "obligations": [
                            {
                                "obligation_id": "obl.loan.disclosure",
                                "source_id": source_id,
                                "source_section": section,
                                "statement": "Loan terms must be disclosed before approval.",
                                "concepts": ["loan", "disclosure"],
                                "applicability_expression": "business_type includes personal_loan",
                                "required_surfaces": ["frontend_h5"],
                                "source_refs": [
                                    {
                                        "source_id": source_id,
                                        "source_section": section,
                                    }
                                ],
                            }
                        ],
                    }
                )
            )
        obligation = payload["obligations"][0]
        return ModelResponse(
            content=json.dumps(
                {
                    "contract": "control_draft_set.v1",
                    "version": "1.0",
                    "status": "draft",
                    "controls": [
                        {
                            "control_id": "loan_disclosure.before_approval",
                            "module_id": "loan_disclosure",
                            "title": "Loan terms are disclosed before approval",
                            "severity": "high",
                            "obligation_ids": [obligation["obligation_id"]],
                            "evidence_requirements": [
                                {
                                    "surface": "frontend_h5",
                                    "minimum_strength": "static_proof",
                                    "rationale": (
                                        "The user-facing disclosure must be visible in the app."
                                    ),
                                }
                            ],
                            "missing_evidence_policy": "block",
                            "reuse_invalidation_keys": ["control_version", "frontend_h5_revision"],
                        }
                    ],
                }
            )
        )

    return StaticModelProvider(response)


def test_phase2_compiles_traceable_controls_with_two_structured_calls(tmp_path: Path) -> None:
    source = tmp_path / "policy.md"
    source.write_text("# Loan disclosure\n\nTerms must be disclosed.\n", encoding="utf-8")
    provider = _provider_for_valid_compilation()
    result = Phase2CompilationService(tmp_path / "workspace", provider).compile(
        [source], {source.as_posix(): "google_play"}
    )

    assert len(provider.requests) == 2
    assert all(not request.tools for request in provider.requests)
    assert result.control_validation.valid is True
    assert (tmp_path / "workspace" / "setup" / "sources.json").is_file()
    assert (tmp_path / "workspace" / "setup" / "obligations.json").is_file()
    assert (tmp_path / "workspace" / "setup" / "controls_draft.json").is_file()
    assert (tmp_path / "workspace" / "setup" / "controls.json").is_file()
    control = result.controls.controls[0]
    assert control.obligation_ids == ["obl.loan.disclosure"]
    assert control.source_refs[0].source_id == result.source_registry.sources[0].source_id


def test_control_claim_routes_are_policy_candidates_not_required_surfaces() -> None:
    source_ref = SourceRef(source_id="policy-1", source_section="section-1")
    route = EvidenceProofRoute(
        route_id="android-disclosure",
        surface="android_native",
        claim_to_prove="A native disclosure entry exists.",
        expected_evidence_strength="static_proof",
        why_this_surface="The app may deliver the disclosure through native UI.",
        acceptance_criteria=[
            AcceptanceCriterion(
                criterion_id="native-entry",
                criterion_type="presence",
                statement="A native disclosure entry is present.",
                scope="Production Android UI code.",
            )
        ],
        proof_limits=["Static code cannot prove the dialog was shown at runtime."],
    )
    claim = EvidenceClaim(
        claim_id="disclosure-entry",
        statement="The app provides a disclosure entry before the relevant action.",
        obligation_ids=["obl.loan.disclosure"],
        source_refs=[source_ref],
        proof_routes=[route],
    )
    control = ControlDraft(
        control_id="loan.disclosure",
        module_id="loan",
        title="Disclosure entry",
        severity="high",
        obligation_ids=["obl.loan.disclosure"],
        source_refs=[source_ref],
        applicability_condition={"kind": "unknown", "reason": "profile-dependent"},
        candidate_surfaces=["android_native"],
        required_surfaces=["android_native"],
        evidence_requirements={
            "android_native": EvidenceRequirement(
                minimum_strength="static_proof",
                rationale="The selected route can prove the static entry.",
                obligation_ids=["obl.loan.disclosure"],
                source_refs=[source_ref],
            )
        },
        missing_evidence_policy="block",
        reuse_invalidation_keys=["control_version"],
        evidence_claims=[claim],
    )

    assert control.evidence_claims[0].proof_routes[0].surface == "android_native"
    assert control.candidate_surfaces == ["android_native"]


def test_phase2_accepts_mixed_obligation_and_no_obligation_sections(tmp_path: Path) -> None:
    source = tmp_path / "mixed-policy.md"
    source.write_text(
        "# Normative\n\n"
        + "Terms must be disclosed before approval. " * 8
        + "\n\n# Informational\n\n"
        + "This section explains the document structure. " * 8
        + "\n",
        encoding="utf-8",
    )

    def response(request: object) -> ModelResponse:
        kind = request.request_kind  # type: ignore[attr-defined]
        payload = json.loads(request.messages[1]["content"])  # type: ignore[attr-defined]
        if kind == "obligation_extraction":
            sections = payload["sources"][0]["sections"]
            decisions = []
            obligations = []
            for section in sections:
                if section["title"] == "Normative":
                    decisions.append(
                        {
                            "section_id": section["section_id"],
                            "decision": "obligations_extracted",
                            "obligation_ids": ["obl.mixed.disclosure"],
                        }
                    )
                    obligations.append(
                        {
                            "obligation_id": "obl.mixed.disclosure",
                            "source_id": payload["source_id"],
                            "source_section": section["section_id"],
                            "statement": "Terms must be disclosed before approval.",
                            "concepts": ["loan", "disclosure"],
                            "applicability_expression": "business_type includes personal_loan",
                            "required_surfaces": ["frontend_h5"],
                            "source_refs": [
                                {
                                    "source_id": payload["source_id"],
                                    "source_section": section["section_id"],
                                }
                            ],
                        }
                    )
                else:
                    decisions.append(
                        {
                            "section_id": section["section_id"],
                            "decision": "no_obligation",
                            "reason": "Informational text only.",
                        }
                    )
            return ModelResponse(
                content=json.dumps(
                    {
                        "contract": "obligation_extraction_batch.v1",
                        "version": "1.0",
                        "source_id": payload["source_id"],
                        "batch_id": payload["batch_id"],
                        "section_decisions": decisions,
                        "obligations": obligations,
                    }
                )
            )

        obligation = payload["obligations"][0]
        return ModelResponse(
            content=json.dumps(
                {
                    "contract": "control_draft_set.v1",
                    "version": "1.0",
                    "status": "draft",
                    "controls": [
                        {
                            "control_id": "mixed.disclosure",
                            "module_id": "loan_disclosure",
                            "title": "Loan terms are disclosed before approval",
                            "severity": "high",
                            "obligation_ids": [obligation["obligation_id"]],
                            "evidence_requirements": [
                                {
                                    "surface": "frontend_h5",
                                    "minimum_strength": "static_proof",
                                    "rationale": "Disclosure is user-facing.",
                                }
                            ],
                            "missing_evidence_policy": "block",
                            "reuse_invalidation_keys": ["control_version"],
                        }
                    ],
                }
            )
        )

    result = Phase2CompilationService(
        tmp_path / "workspace",
        StaticModelProvider(response),
        source_builder=SourceRegistryBuilder(max_section_chars=500),
    ).compile([source])

    assert len(result.source_registry.sources[0].sections) == 2
    assert len(result.obligations.obligations) == 1
    assert result.control_validation.valid is True


def test_invalid_control_draft_does_not_write_validated_controls(tmp_path: Path) -> None:
    source = tmp_path / "policy.txt"
    source.write_text("Terms must be disclosed.", encoding="utf-8")

    def invalid_response(request: object) -> ModelResponse:
        kind = request.request_kind  # type: ignore[attr-defined]
        payload = json.loads(request.messages[1]["content"])  # type: ignore[attr-defined]
        if kind == "obligation_extraction":
            return ModelResponse(
                content=json.dumps(
                    {
                        "contract": "obligation_extraction_batch.v1",
                        "version": "1.0",
                        "source_id": payload["source_id"],
                        "batch_id": payload["batch_id"],
                        "section_decisions": [],
                        "obligations": [],
                    }
                )
            )
        return ModelResponse(
            content=json.dumps(
                {
                    "contract": "control_draft_set.v1",
                    "version": "1.0",
                    "status": "draft",
                    "controls": [],
                }
            )
        )

    invalid_provider = StaticModelProvider(invalid_response)
    with pytest.raises(Phase2CompilationError):
        Phase2CompilationService(tmp_path / "workspace", invalid_provider).compile([source])
    assert not (tmp_path / "workspace" / "setup" / "controls.json").exists()
    assert (tmp_path / "workspace" / "setup" / "sources.json").is_file()


def test_validator_separates_section_coverage_from_control_provenance() -> None:
    source = ComplianceSource(
        source_id="policy",
        path="policy.md",
        title="Policy",
        sha256="a" * 64,
        source_family="regulator",
        media_type="md",
        extraction_status="ok",
        sections=[
            SourceSection(section_id="one", title="One", text="One", ordinal=1),
            SourceSection(section_id="two", title="Two", text="Two", ordinal=2),
        ],
    )
    obligation = Obligation(
        obligation_id="obl.one",
        source_id="policy",
        source_section="one",
        statement="One applies.",
        concepts=["one"],
        applicability_expression="business_type includes personal_loan",
        required_surfaces=["frontend_h5"],
        source_refs=[SourceRef(source_id="policy", source_section="one")],
    )
    draft = ControlDraft(
        control_id="control.one",
        module_id="safe.module",
        title="One",
        severity="low",
        obligation_ids=["obl.one"],
        source_refs=[SourceRef(source_id="policy", source_section="one")],
        applicability_expression="business_type includes personal_loan and self_lending == true",
        required_surfaces=["frontend_h5"],
        evidence_requirements={
            "frontend_h5": EvidenceRequirement(minimum_strength="static_proof", rationale="fixture")
        },
        missing_evidence_policy="block",
        reuse_invalidation_keys=["control_version"],
    )
    validation = ControlValidator().validate(
        SourceRegistry(sources=[source], version="1.0"),
        ObligationSet(version="1.0", status="draft", obligations=[obligation]),
        ControlDraftSet(version="1.0", controls=[draft]),
    )

    assert validation.valid is False
    assert not any(
        "not covered by an obligation: policy/two" in error for error in validation.errors
    )
    assert any("narrows applicability" in error for error in validation.errors)


def test_validator_accepts_empty_validated_set_when_all_sections_have_no_obligation() -> None:
    source = ComplianceSource(
        source_id="informational",
        path="informational.md",
        title="Informational",
        sha256="c" * 64,
        source_family="internal",
        media_type="md",
        extraction_status="ok",
        sections=[SourceSection(section_id="one", title="One", text="One", ordinal=1)],
    )

    validation = ControlValidator().validate(
        SourceRegistry(sources=[source], version="1.0"),
        ObligationSet(version="1.0", status="draft", obligations=[]),
        ControlDraftSet(version="1.0", controls=[]),
    )

    assert validation.valid is True


def test_validator_rejects_duplicate_obligations_or_orphans() -> None:
    source = ComplianceSource(
        source_id="policy",
        path="policy.md",
        title="Policy",
        sha256="b" * 64,
        source_family="regulator",
        media_type="md",
        extraction_status="ok",
        sections=[SourceSection(section_id="one", title="One", text="One", ordinal=1)],
    )
    obligation = Obligation(
        obligation_id="obl.same",
        source_id="policy",
        source_section="one",
        statement="One applies.",
        concepts=["one"],
        applicability_expression="business_type includes personal_loan",
        required_surfaces=["frontend_h5"],
        source_refs=[SourceRef(source_id="policy", source_section="one")],
    )
    validation = ControlValidator().validate(
        SourceRegistry(sources=[source], version="1.0"),
        ObligationSet(version="1.0", status="draft", obligations=[obligation, obligation]),
        ControlDraftSet(version="1.0", controls=[]),
    )

    assert validation.valid is False
    assert validation.duplicate_obligation_ids == ["obl.same"]
    assert any("duplicate obligation_id" in error for error in validation.errors)


def test_applicability_validator_rejects_unbounded_expression() -> None:
    errors = validate_applicability_expression("python eval os.system")
    assert errors
